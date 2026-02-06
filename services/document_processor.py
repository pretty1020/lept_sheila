"""
LEPT AI Reviewer - Document Processing Service
Enhanced to handle various file formats and sizes
"""

import io
from typing import Optional, Tuple
from pathlib import Path

import streamlit as st


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[bool, str]:
    """
    Extract text from a PDF file using multiple methods.
    
    Args:
        file_bytes: PDF file as bytes
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    text_parts = []
    
    # Method 1: Try PyPDF2
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            except Exception:
                continue
        
    except ImportError:
        pass
    except Exception as e:
        print(f"PyPDF2 extraction error: {e}")
    
    # Method 2: Try pdfplumber if PyPDF2 didn't get much text
    if len("".join(text_parts)) < 100:
        try:
            import pdfplumber
            
            pdf_file = io.BytesIO(file_bytes)
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text)
                    except Exception:
                        continue
        except ImportError:
            pass
        except Exception as e:
            print(f"pdfplumber extraction error: {e}")
    
    # Method 3: Try pypdf if others failed
    if len("".join(text_parts)) < 100:
        try:
            from pypdf import PdfReader as PyPdfReader
            
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPdfReader(pdf_file)
            
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                except Exception:
                    continue
        except ImportError:
            pass
        except Exception as e:
            print(f"pypdf extraction error: {e}")
    
    if not text_parts:
        # Return success with a placeholder - allow upload even without text
        return True, "[PDF document uploaded - text extraction limited. The document can still be used for reference.]"
    
    full_text = "\n\n".join(text_parts)
    
    # Clean up the text
    full_text = clean_extracted_text(full_text)
    
    # Don't reject based on text length - allow any content
    if not full_text.strip():
        return True, "[PDF document uploaded - contains minimal extractable text.]"
    
    return True, full_text


def extract_text_from_docx(file_bytes: bytes) -> Tuple[bool, str]:
    """
    Extract text from a DOCX file with multiple fallback methods.
    
    Args:
        file_bytes: DOCX file as bytes
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    text_parts = []
    
    # Method 1: Try python-docx
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text and text.strip():
                text_parts.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        
        # Extract from headers/footers
        for section in doc.sections:
            try:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text and para.text.strip():
                            text_parts.append(para.text)
            except Exception:
                pass
                
    except ImportError:
        return False, "python-docx library not installed. Please contact administrator."
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        # Try alternative method
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            docx_file = io.BytesIO(file_bytes)
            with zipfile.ZipFile(docx_file) as z:
                # Read the main document
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ET.fromstring(xml_content)
                    
                    # Extract all text
                    for elem in tree.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text)
        except Exception as e2:
            print(f"Alternative DOCX extraction error: {e2}")
    
    if not text_parts:
        # Return success with placeholder - allow upload even without extracted text
        return True, "[DOCX document uploaded - text extraction limited. The document can still be used for reference.]"
    
    full_text = "\n\n".join(text_parts)
    
    # Clean up the text
    full_text = clean_extracted_text(full_text)
    
    # Don't reject based on text length - allow any content
    if not full_text.strip():
        return True, "[DOCX document uploaded - contains minimal extractable text.]"
    
    return True, full_text


def extract_text_from_file(uploaded_file) -> Tuple[bool, str]:
    """
    Extract text from an uploaded file (PDF or DOCX).
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Tuple of (success, extracted_text_or_error)
    """
    if uploaded_file is None:
        return False, "No file provided"
    
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    
    # Reset file pointer for potential reuse
    uploaded_file.seek(0)
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        return False, "Unsupported file format. Please upload a PDF or DOCX file."


def clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Strip whitespace
        line = line.strip()
        
        # Skip empty lines (but keep some for paragraph separation)
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        
        # Remove lines that are just special characters
        if all(c in '.-_=*#' for c in line.replace(' ', '')):
            continue
        
        cleaned_lines.append(line)
    
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Remove multiple consecutive newlines (keep max 2)
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text.strip()


def truncate_text_for_ai(text: str, max_chars: int = 15000) -> str:
    """
    Truncate text to fit within AI context limits.
    
    Args:
        text: Full text
        max_chars: Maximum character count
    
    Returns:
        Truncated text
    """
    if len(text) <= max_chars:
        return text
    
    # Try to truncate at a sentence boundary
    truncated = text[:max_chars]
    last_period = truncated.rfind('.')
    
    if last_period > max_chars * 0.8:
        return truncated[:last_period + 1]
    
    return truncated + "..."


def get_text_stats(text: str) -> dict:
    """
    Get statistics about extracted text.
    
    Args:
        text: Extracted text
    
    Returns:
        Dictionary with text statistics
    """
    if not text:
        return {
            "char_count": 0,
            "word_count": 0,
            "line_count": 0,
            "paragraph_count": 0
        }
    
    words = text.split()
    lines = text.split('\n')
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    
    return {
        "char_count": len(text),
        "word_count": len(words),
        "line_count": len(lines),
        "paragraph_count": len(paragraphs)
    }
